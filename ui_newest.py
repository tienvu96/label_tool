"""
Author: tienva@bidv.com.vn
Version: 1
Date: 12/09/2024
"""

import sys
import json
from pathlib import Path
from openpyxl import load_workbook
from PyQt5.QtWidgets import (
    QApplication, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog,
    QMessageBox, QDesktopWidget, QMenu
)
from main_newest import (
    scan_file, label_docx_file, label_xlsx_file_footer,
    classify_document_with_multiple_rules, define_rules, is_file_locked
)
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER


class LabelingApp(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle('Label tool')
        self.setGeometry(100, 100, 600, 400)
        self.center()

        # Main layout
        main_layout = QVBoxLayout()

        # File info layout (file selection)
        file_info_layout = QHBoxLayout()

        # Label for file path
        self.file_path_label = QLabel('File path')
        file_info_layout.addWidget(self.file_path_label)

        # Text field to display the selected file path
        self.file_path_edit = QLineEdit()
        file_info_layout.addWidget(self.file_path_edit)

        # Browse button to select a file
        self.browse_button = QPushButton('Browse')
        self.browse_button.clicked.connect(self.browse_file)
        file_info_layout.addWidget(self.browse_button)

        main_layout.addLayout(file_info_layout)

        # File operations layout (buttons for scanning and labeling)
        file_ops_layout = QHBoxLayout()

        # Button to scan the file
        self.scan_button = QPushButton('Scan')
        self.scan_button.clicked.connect(self.handle_scan_file)
        file_ops_layout.addWidget(self.scan_button)

        # Button to label the file
        self.label_button = QPushButton('Label file')
        self.label_button.clicked.connect(self.handle_label_file)
        file_ops_layout.addWidget(self.label_button)

        # Button for editing label (with options Bí mật, Nội bộ, Công khai)
        self.edit_label_button = QPushButton('Edit Label')
        self.edit_label_menu = QMenu(self)

        # Add options to the Edit Label menu
        self.edit_label_menu.addAction(
            "Bí mật", lambda: self.set_label("Bí mật"))
        self.edit_label_menu.addAction(
            "Nội bộ", lambda: self.set_label("Nội bộ"))
        self.edit_label_menu.addAction(
            "Công khai", lambda: self.set_label("Công khai"))

        # Connect the button to show the menu when clicked
        self.edit_label_button.setMenu(self.edit_label_menu)
        file_ops_layout.addWidget(self.edit_label_button)

        # Button to show file label info
        self.file_info_button = QPushButton('Labeled file information')
        self.file_info_button.clicked.connect(self.handle_file_info)
        file_ops_layout.addWidget(self.file_info_button)

        main_layout.addLayout(file_ops_layout)

        # Text area to display file details or scan results
        self.file_details_text = QTextEdit()
        main_layout.addWidget(self.file_details_text)

        # Set the main layout for the window
        self.setLayout(main_layout)

    def center(self):
        """Centers the window on the screen"""
        screen = QDesktopWidget().screenGeometry()
        window = self.frameGeometry()
        window.moveCenter(screen.center())
        self.move(window.topLeft())

    def browse_file(self):
        """Opens a file dialog to browse and select a file"""
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(
            self, "Chọn file", "", "All Files (*)"
        )
        if file_path:
            self.file_path_edit.setText(file_path)

    def handle_scan_file(self):
        """Handles the scan operation and displays the results."""
        file_path = self.file_path_edit.text()
        if not file_path:
            self.file_details_text.setText(
                "Vui lòng chọn một tệp trước khi scan.")
            return

        if is_file_locked(file_path):
            self.file_details_text.setText(
                "File đang mở. Vui lòng đóng file trước khi quét tệp."
            )
            return

        result_scan = scan_file(file_path)
        rule = define_rules()
        classify, sms = classify_document_with_multiple_rules(
            result_scan, rule)

        if classify == "Confidential":
            self.file_details_text.setText(sms)
        elif classify == "Internal":
            self.file_details_text.setText(
                "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!")
        elif classify == "chưa làm" or classify == "Unsupport file":
            self.file_details_text.setText("Hiện tại, không hỗ trợ file.")


    def handle_label_file(self):
        """Handles the labeling operation for DOCX and XLSX files."""
        file_path = self.file_path_edit.text()

        if not file_path:
            self.file_details_text.setText(
                "Vui lòng chọn một tệp trước khi label.")
            return

        if is_file_locked(file_path):
            self.file_details_text.setText(
                "File đang mở. Vui lòng đóng file trước khi gán nhãn."
            )
            return

        if file_path.endswith('.docx'):
            results = scan_file(file_path)
            rules = define_rules()
            category, sms = classify_document_with_multiple_rules(results, rules)
            if category:
                label_text = label_docx_file(file_path, category)
                self.file_details_text.setText(
                    f"Đã gán nhãn '{label_text}' cho tài liệu.")
                self.show_message(
                    f"Đã gán nhãn '{label_text}' cho file thành công.")

        elif file_path.endswith(".xlsx"):
            results = scan_file(file_path)
            rules = define_rules()
            category, sms = classify_document_with_multiple_rules(results, rules)

            if category:
                label_text = label_xlsx_file_footer(file_path, category)
                self.file_details_text.setText(
                    f"Đã gán nhãn '{label_text}' cho tài liệu.")
                self.show_message(
                    f"Đã gán nhãn '{label_text}' cho file thành công.")

        else:
            self.file_details_text.setText(
                "Chức năng gán nhãn chỉ hỗ trợ định dạng DOCX, XLSX."
            )

    def set_label(self, label_type):
        """Set the label type based on user selection and ask for confirmation"""
        msg_box = QMessageBox()
        msg_box.setIcon(QMessageBox.Question)
        msg_box.setWindowTitle("Xác nhận")
        msg_box.setText(
            f"Bạn có chắc chắn muốn gán nhãn '{label_type}' cho file không?")
        msg_box.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        result = msg_box.exec_()

        if result == QMessageBox.Yes:
            self.file_details_text.setText(f"Đang gán nhãn: {label_type}")
            self.label_file_with_new_label(label_type)

    def label_file_with_new_label(self, label_type):
        """Label the file with the new label selected by the user"""
        file_path = self.file_path_edit.text()

        if not file_path:
            self.file_details_text.setText(
                "Vui lòng chọn một tệp trước khi label.")
            return

        if is_file_locked(file_path):
            self.file_details_text.setText(
                "File đang mở. Vui lòng đóng file trước khi gán nhãn."
            )
            return

        if file_path.endswith('.docx'):
            label_text = self.edit_label_docx_file(file_path, label_type)
            self.file_details_text.setText(
                f"Đã gán nhãn '{label_text}' cho file DOCX.")
        elif file_path.endswith('.xlsx'):
            label_text = self.edit_label_xlsx_file_footer(
                file_path, label_type)
            self.file_details_text.setText(
                f"Đã gán nhãn '{label_text}' cho file XLSX.")
        else:
            self.file_details_text.setText(
                "Chỉ hỗ trợ gán nhãn cho file DOCX và XLSX."
            )

    def edit_label_docx_file(self, file_path, label_type):
        """Add or overwrite a label in the footer of a DOCX file based on user selection."""
        document = Document(file_path)

        if label_type == "Bí mật":
            label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
        elif label_type == "Nội bộ":
            label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
        else:
            label_text = "Tài liệu Công khai của BIDV."

        section = document.sections[0]
        section.top_margin = Cm(2)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)

        footer = section.footer

        # Remove all existing paragraphs in the footer
        for element in footer._element.xpath('.//w:p'):
            element.getparent().remove(element)

        # Add a new paragraph that contains both the page number and the label text
        paragraph = footer.add_paragraph()

        # Set a tab stop at the right margin for the page number
        right_margin_position = section.page_width - \
            section.right_margin - section.left_margin
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            right_margin_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES
        )

        # Add the label text (left-aligned)
        run_label = paragraph.add_run(label_text)
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(12)

        # Add a tab to move the cursor to the right side of the page
        run_label.add_tab()

        # Add the page number field (aligned to the right tab stop)
        run_page = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')  # Begin field
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')  # Field instruction text
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')  # Separate field
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')  # End field
        fldChar3.set(qn('w:fldCharType'), 'end')

        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)
        run_page._r.append(fldChar3)

        # Set font for the page number
        run_page.font.name = 'Times New Roman'
        run_page.font.size = Pt(12)

        # Align the paragraph to left (default behavior)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Save the document
        document.save(file_path)

        return label_text

    def edit_label_xlsx_file_footer(self, file_path, label_type):
        """Add or overwrite a label in the footer of an XLSX file with label on the left and page number on the right."""
        workbook = load_workbook(file_path)
        
        # Duyệt qua tất cả các sheet trong file Excel
        for sheet in workbook.worksheets:
            # Xác định nhãn dựa trên loại nhãn
            if label_type == "Bí mật":
                label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
            elif label_type == "Nội bộ":
                label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
            else:
                label_text = "Tài liệu Công khai của BIDV."

            # Cài đặt footer với nhãn ở bên trái và số trang ở bên phải
            footer_text = f"&L{label_text}&RTrang &P"

            # Gán footer cho sheet
            sheet.oddFooter.center.text = footer_text
            sheet.oddFooter.center.size = 12  # Cỡ chữ 12
            sheet.oddFooter.center.font = "Times New Roman"  # Font Times New Roman

        # Lưu lại workbook với các thay đổi
        workbook.save(file_path)

        return label_text

    def handle_file_info(self):
        """Handles displaying the file path and label text information"""
        file_path = self.file_path_edit.text()
        if not file_path:
            self.file_details_text.setText(
                "Vui lòng chọn một tệp trước khi hiển thị thông tin.")
            return

        if is_file_locked(file_path):
            self.file_details_text.setText(
                "File đang mở. Vui lòng đóng file trước khi xem thông tin tệp."
            )
            return

        results, message = scan_file(file_path)
        rules = define_rules()
        category = classify_document_with_multiple_rules(results, rules)

        if file_path.endswith('.docx'):
            label_text = label_docx_file(file_path, category)
        elif file_path.endswith('.xlsx'):
            label_text = label_xlsx_file_footer(file_path, category)
        else:
            label_text = "Hiện tại chỉ hỗ trợ file DOCX, XLSX, CSV."

        result_info = {"Path": file_path}
        file_info = json.dumps(result_info, indent=4, ensure_ascii=False)
        self.file_details_text.setText(file_info)
        return file_info

    def show_message(self, message):
        """Show a message box with the provided message."""
        msg_box = QMessageBox()
        msg_box.setIcon(QMessageBox.Information)
        msg_box.setText(message)
        msg_box.setWindowTitle("Thông báo")
        msg_box.exec_()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = LabelingApp()
    ex.show()
    sys.exit(app.exec_())
