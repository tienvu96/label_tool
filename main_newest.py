import pandas as pd
import json
import io
import csv
import re
from openpyxl import load_workbook
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.shared import Pt
from pathlib import Path
from docx import Document
import io
import csv
import pandas as pd
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
# Function to label DOCX file
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER


def is_file_locked(file_path):
    """Check if the file is currently locked (opened by another application)."""
    try:
        # Try opening the file in append mode (or any mode) to check for a lock
        with open(file_path, 'a'):
            pass
        return False  # File is not locked
    except IOError:
        return True  # File is locked


# Function to read the JSON file and load the keywords
# Load keywords from JSON file
def load_keywords_from_json():
    """Load keywords from a JSON file."""
    json_file = "/Users/vuanhtien/Documents/CodeForMe/BIDV/label_tool/keyword.json"
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    return data


def extract_and_iterate_docx_content(file_path, table_id=None, **pandas_kwargs):
    """Extract all content from a DOCX file, including headers, footers, paragraphs, and tables."""
    document = Document(file_path)

    def extract_single_table(table):
        memory_file = io.StringIO()
        csv_writer = csv.writer(memory_file)
        for row in table.rows:
            csv_writer.writerow([cell.text.strip() for cell in row.cells])
        memory_file.seek(0)
        return pd.read_csv(memory_file, **pandas_kwargs)

    content_list = []

    # Extract paragraphs and tables from the body
    for child_element in document.element.body.iterchildren():
        if isinstance(child_element, CT_P):
            # Convert to lowercase
            paragraph = Paragraph(child_element, document).text.strip().lower()
            if paragraph:  # Only append non-empty paragraphs
                content_list.append(paragraph)
        elif isinstance(child_element, CT_Tbl):
            table = Table(child_element, document)
            if table_id is None or document.tables.index(table) == table_id:
                # Convert the table to string before adding to the content_list
                table_df = extract_single_table(table)
                # Convert DataFrame to string in lowercase
                content_list.append(table_df.to_string(index=False).lower())

    # Extract headers and footers
    for section in document.sections:
        # Extract headers
        for header in section.header.paragraphs:
            header_text = header.text.strip().lower()  # Convert to lowercase
            if header_text:
                content_list.append(header_text)

        # Extract footers
        # for footer in section.footer.paragraphs:
        #     footer_text = footer.text.strip().lower()  # Convert to lowercase
        #     if footer_text:
        #         content_list.append(footer_text)

    return content_list


# Unified function to search for both keywords and regex patterns
def find_keywords_and_patterns_in_docx(content_list, keywords_dict, patterns):
    """Search for both keywords and patterns in DOCX content and return the count of occurrences."""
    found_keywords = []
    found_patterns = []

    # Khởi tạo bộ đếm cho từ khóa và pattern
    keyword_counts = {}
    pattern_counts = {}

    # Chuyển nội dung DOCX thành một chuỗi duy nhất và chuyển thành chữ thường
    docx_content_lower = " ".join(
        [str(content).strip().lower() for content in content_list])

    # Tìm từ khóa
    for keyword_list in keywords_dict.values():
        for keyword in keyword_list:
            # Chuyển từ khóa thành chữ thường và loại bỏ khoảng trắng
            keyword_lower = keyword.lower().strip()
            keyword_counts[keyword_lower] = 0  # Khởi tạo bộ đếm cho từ khóa

            # Đếm số lần từ khóa xuất hiện trong nội dung DOCX
            keyword_counts[keyword_lower] += docx_content_lower.count(
                keyword_lower)

            # Nếu từ khóa xuất hiện ít nhất 1 lần, thêm vào kết quả
            if keyword_counts[keyword_lower] > 0:
                found_keywords.append({
                    "Found Keyword": keyword,
                    "num of the same keyword": keyword_counts[keyword_lower]
                })

    # Tìm pattern (mẫu regex)
    for pattern_name, pattern_regex in patterns.items():
        pattern_counts[pattern_name] = 0  # Khởi tạo bộ đếm cho pattern

        # Tìm tất cả các pattern khớp trong nội dung DOCX
        matches = pattern_regex.findall(docx_content_lower)
        # Cộng số lượng khớp với bộ đếm của pattern
        pattern_counts[pattern_name] += len(matches)

        # Nếu pattern xuất hiện ít nhất 1 lần, thêm vào kết quả
        if pattern_counts[pattern_name] > 0:
            found_patterns.append({
                "Pattern Name": pattern_name,
                "num of the same pattern": pattern_counts[pattern_name]
            })

    # Chuyển đổi kết quả sang JSON và lưu vào biến
    result = {
        "Keywords": found_keywords,
        "Patterns": found_patterns
    }

    json_result = json.dumps(result, indent=4, ensure_ascii=False)

    # Trả về biến chứa kết quả JSON
    return json_result


# Main function to check keywords and patterns in a DOCX file
def check_keywords_and_patterns_in_docx(file_path, patterns):
    """
    Check the DOCX content for both keywords and patterns.
    """
    # Load keywords from JSON file
    keywords_dict = load_keywords_from_json()

    # Extract content from DOCX
    content_list = extract_and_iterate_docx_content(file_path)

    # Find both keywords and patterns
    result = find_keywords_and_patterns_in_docx(
        content_list, keywords_dict, patterns)

    return result


def extract_and_iterate_excel_content(file_path, **pandas_kwargs):
    """Extracts content from all sheets of an Excel file and returns the content, including column headers."""
    # Đảm bảo pandas đọc tất cả các giá trị dưới dạng chuỗi để không mất số 0 đầu
    # dtype=str buộc đọc tất cả dữ liệu dưới dạng chuỗi
    sheet_dict = pd.read_excel(
        file_path, sheet_name=None, dtype=str, **pandas_kwargs)

    content_dict = {}

    for sheet_name, df in sheet_dict.items():
        content_list = []

        # Thêm tên cột vào content_list, bỏ qua các cột chứa 'Unnamed'
        headers = df.columns
        content_list.extend([str(header).strip()
                            for header in headers if "Unnamed" not in header])

        # Thêm nội dung của từng ô vào content_list, bỏ qua các ô chứa 'Unnamed'
        for _, row in df.iterrows():
            for cell in row:
                cell_str = str(cell).strip()
                if pd.notnull(cell) and "Unnamed" not in cell_str:
                    content_list.append(cell_str)

        content_dict[sheet_name] = content_list

    return content_dict


file_path = "/Users/vuanhtien/Documents/CodeForMe/BIDV/label_tool/File test 092024/202409 Testcase thong tin dinh danh.xlsx"
excel_content = extract_and_iterate_excel_content(file_path)
# print("a" *100)
# print(excel_content)


def find_keywords_and_patterns_in_excel(excel_content_dict, keywords_dict, patterns):
    """Search for both keywords and patterns in Excel content and return the count of occurrences in each sheet."""
    found_keywords = []
    found_patterns = []

    for sheet_name, excel_content in excel_content_dict.items():
        keyword_counts = {}
        pattern_counts = {}

        # Tìm từ khóa
        for keyword_list in keywords_dict.values():
            for keyword in keyword_list:
                # Chuyển từ khóa thành chữ thường và loại bỏ khoảng trắng
                keyword_lower = keyword.lower().strip()
                # Khởi tạo bộ đếm cho từ khóa
                keyword_counts[keyword_lower] = 0

                # Duyệt qua nội dung Excel và đếm số lần từ khóa xuất hiện
                for content in excel_content:
                    # Chuyển nội dung thành chuỗi, loại bỏ khoảng trắng và chuyển chữ thường
                    content_clean = str(content).strip().lower()
                    # Đếm số lần từ khóa xuất hiện
                    keyword_counts[keyword_lower] += content_clean.count(
                        keyword_lower)

                # Nếu từ khóa xuất hiện ít nhất 1 lần, lưu kết quả vào danh sách
                if keyword_counts[keyword_lower] > 0:
                    found_keywords.append({
                        "Sheet": sheet_name,
                        "Found Keyword": keyword,
                        "num of the same keyword": keyword_counts[keyword_lower]
                    })

        # Tìm pattern (mẫu regex)
        for pattern_name, pattern_regex in patterns.items():
            pattern_counts[pattern_name] = 0  # Khởi tạo bộ đếm cho pattern

            # Duyệt qua từng nội dung trong Excel và tìm pattern
            for content in excel_content:
                # Tìm tất cả các pattern khớp
                matches = pattern_regex.findall(content)
                # Cộng số lượng khớp với bộ đếm của pattern
                pattern_counts[pattern_name] += len(matches)

            # Nếu pattern xuất hiện ít nhất 1 lần, thêm kết quả vào danh sách
            if pattern_counts[pattern_name] > 0:
                found_patterns.append({
                    "Sheet": sheet_name,
                    "Pattern Name": pattern_name,
                    "num of the same pattern": pattern_counts[pattern_name]
                })

    # Chuyển đổi kết quả sang JSON và lưu vào biến
    result = {
        "Keywords": found_keywords,
        "Patterns": found_patterns
    }

    json_result = json.dumps(result, indent=4, ensure_ascii=False)

    # Trả về biến chứa kết quả JSON
    return json_result


# Main function to check keywords and patterns in an Excel file
def check_keywords_and_patterns_in_excel(xlsx_file, patterns):
    """Check an Excel file for both keywords and patterns."""
    # Load keywords từ JSON
    keywords_dict = load_keywords_from_json()

    # Trích xuất nội dung từ file Excel
    excel_content_dict = extract_and_iterate_excel_content(xlsx_file)

    # Find both keywords and patterns
    result = find_keywords_and_patterns_in_excel(
        excel_content_dict, keywords_dict, patterns)

    return result


def define_rules():
    rules = {
        "rule_1": {
            "email": "email",
            "id number (cccd/cmnd)": "id number (cccd/cmnd)",
            "tên khách hàng": ["tên khách hàng", "họ tên", "họ và tên", "họ tên", "name", "full name"],
            "địa chỉ": "địa chỉ",
            "số điện thoại": "số điện thoại"
        },
        "rule_2": {
            "số thẻ": "số thẻ",
            "cvv": "cvv",
            "ngày hết hạn": "ngày hết hạn",
            "tên chủ thẻ": ["tên khách hàng", "chủ thẻ", "tên chủ thẻ", "họ và tên", "họ tên", "name", "full name"]
        },
        "rule_3": {
            "email": "email",
            "số điện thoại": "số điện thoại",
            "số tài khoản ẩn": "số tài khoản ẩn",
            "giá trị tiền tối thiểu 6 chữ số": "giá trị tiền tối thiểu 6 chữ số",
            "as1111": "sf1111",
            "as23333": "sf2222",
            "as555": "s555f",
            "as9999": "s555666f"
        },
        "rule_4": ["mật", "tuyệt mật"]
    }

    return rules


def scan_file(file_path):
    """Scan a file and check for both keywords and patterns depending on the file type."""
    if not file_path:
        return None, "Vui lòng chọn một file trước khi scan."

    file_extension = Path(file_path).suffix.lower()

    # Define patterns (shortened for brevity, include all necessary patterns)
    patterns = {
        "stk 10 số chứa 3 BDS": re.compile(r'\b(111|116|117|118|119|120|121|122|123|124|125|126|128|129|130|131|132|133|134|135|136|138|139|140|141|144|145|147|149|150|151|159|160|166|168|169|177|180|181|186|188|189|199|211|212|213|214|215|216|217|220|222|256|260|261|268|279|289|310|311|313|314|315|317|318|319|321|328|330|341|345|351|362|368|371|375|376|390|395|398|411|421|425|426|427|428|431|432|433|440|441|443|444|448|450|451|452|455|460|461|465|466|468|471|480|482|483|486|488|501|502|505|512|513|518|520|522|531|532|540|556|558|560|561|562|566|565|570|573|580|581|590|601|602|611|615|620|621|625|631|632|633|635|636|641|642|646|650|651|652|653|655|656|661|670|671|672|679|680|686|691|696|701|702|710|711|721|729|730|735|737|741|742|745|748|750|753|760|761|762|766|780|785|788)\d{7}\b'),
        "cvv": re.compile(r"(?i)(\bCVV\b[\s\S]*?\b\d{3}\b)"),
        "id number (cccd/cmnd)": re.compile(r'\b(?:cmnd|cccd|CMND|CCCD)\b.*\b\d{9}\b|\b(?:cmnd|cccd|CMND|CCCD)\b.*\b\d{12}\b|\b\d{9}\b|\b\d{12}\b'),
        "địa chỉ": re.compile(r'(\d+)\s+(đường|phố|phường|quận|thành phố)\s+(\w+),\s+(\w+),\s+(\w+)'),
        "giá trị tiền tối thiểu 6 chữ số": re.compile(r'\b\d{1,3}(,\d{3}){1,}\b'),
        "số điện thoại": re.compile(r'\b(0\d{9}|\+[\d]{11})\b'),
        "email": re.compile(r'\S+@\S+'),
        "số thẻ": re.compile(r'''
                                \b                                        # Bắt đầu tại ranh giới từ
                                (?:9704|476632|411153|428695|427126|402460|406220|511957|517107|517453|542726|530515|515110|51511)
                                (                                         # Bắt đầu nhóm lựa chọn số
                                    \d{10,15}                             # 10 đến 15 số liên tiếp (cho cả số tài khoản và số thẻ)
                                    |\d{16}                               # Bắt chuỗi 16 số liên tiếp
                                    |\d{19}                               # Bắt chuỗi 19 số liên tiếp
                                    |\d{20}                               # Bắt chuỗi 20 số liên tiếp
                                    |\d{3} \d{4} \d{4} \d{4} \d{4}        # 16 số chia thành các nhóm 3-4-4-4-4
                                    |\d{4} \d{4} \d{4} \d{4} \d{3}        # 19 số chia thành nhóm 4-4-4-4-3
                                    |\d{4} \d{4} \d{4} \d{4} \d{4}        # 20 số chia thành nhóm 4-4-4-4-4
                                )
                                \b                                        # Kết thúc tại ranh giới từ
                            ''', re.VERBOSE),
        "số tài khoản ẩn": re.compile(r'\b(?:9704|476632|411153|428695|427126|402460|406220|511957|517107|517453|542726|530515|515110)(?:\d{2}xxxx\d{4}|\d{3}xxxx\d{4}|\d{4}xxxx\d{4})\b')
    }

    # Initialize results variable
    results = None

    # Scan depending on file extension
    if file_extension == '.docx':
        results = check_keywords_and_patterns_in_docx(file_path, patterns)
    elif file_extension == '.xlsx':
        results = check_keywords_and_patterns_in_excel(file_path, patterns)
    elif file_extension in [".pdf", ".ptpx", "tsv", "txt", "py", "png", "jpg"]:
        results = "chưa làm"

    # Ensure results is not None before parsing
    if results is None:
        return None, "Không có kết quả hợp lệ từ file đã chọn."

    results_json = json.loads(results)
    # Return results as a single dictionary
    return results_json


def convert_string_to_json_rule_one(input_string):
    """
    Hàm chuyển đổi chuỗi đầu vào thành JSON với các thông tin từ khóa khớp và số lượng pattern.
    """
    # Tách chuỗi để lấy các thông tin cần thiết
    matched_part = input_string.split(
        "matched: ")[1].split(". Pattern counts: ")[0]
    patterns_part = input_string.split("Pattern counts: ")[1]

    # Chuyển phần matched thành danh sách các keys
    matched_keys = matched_part.split(", ")

    # Chuyển phần pattern counts thành từ điển
    pattern_counts = {}
    for pattern in patterns_part.split(", "):
        key, value = pattern.split(": ")
        pattern_counts[key.strip()] = int(value.strip())

    # Tạo cấu trúc JSON cuối cùng
    result_json = {
        "rule": 1,
        "matched_keys": matched_keys,
        "pattern_counts": pattern_counts
    }

    # Chuyển đổi thành chuỗi JSON đẹp mắt
    pretty_json = json.dumps(result_json, indent=4, ensure_ascii=False)

    # Trả về chuỗi JSON
    return pretty_json


def convert_string_to_json_rule_two(input_string):
    """
    Hàm chuyển đổi chuỗi đầu vào thành JSON với thông tin các từ khóa khớp từ rule.
    """
    # Tách chuỗi để lấy thông tin rule và các từ khóa khớp
    rule_part = input_string.split("rule ")[1].split(":")[0]
    matched_part = input_string.split(": ")[1]

    # Chuyển phần matched thành danh sách các keys
    matched_keys = matched_part.split(", ")

    # Tạo cấu trúc JSON cuối cùng
    result_json = {
        "rule": int(rule_part),  # Chuyển rule thành số nguyên
        "matched_keys": matched_keys
    }

    # Chuyển đổi thành chuỗi JSON đẹp mắt
    pretty_json = json.dumps(result_json, indent=4, ensure_ascii=False)

    # Trả về chuỗi JSON
    return pretty_json


def convert_rule_string_to_json_rule_four(input_string):
    """
    Hàm chuyển đổi chuỗi đầu vào thành JSON với thông tin các từ khóa khớp từ rule.
    """
    # Tách chuỗi để lấy thông tin rule và các từ khóa khớp
    rule_part = input_string.split("rule ")[1].split(":")[0]
    matched_part = input_string.split(": ")[1]

    # Chuyển phần matched thành danh sách các keys (nếu chỉ có 1 từ khóa, vẫn đưa vào danh sách)
    matched_keys = [matched_part]

    # Tạo cấu trúc JSON cuối cùng
    result_json = {
        "rule": int(rule_part),  # Chuyển rule thành số nguyên
        "matched_keys": matched_keys
    }

    # Chuyển đổi thành chuỗi JSON đẹp mắt
    pretty_json = json.dumps(result_json, indent=4, ensure_ascii=False)

    # Trả về chuỗi JSON
    return pretty_json


def classify_document_with_multiple_rules(results, rules):
    # Check if results is "chưa làm" or None
    if results == "chưa làm" or results is None:
        return "chưa làm", "Hiện tại không hỗ trợ định dạng tệp."

    # Ensure results is a dictionary and contains 'Keywords' key
    if not isinstance(results, dict):
        return "Unsupport file", "Kết quả không hợp lệ."

    # Check if 'Keywords' exists and is a list of dictionaries
    if 'Keywords' not in results or not isinstance(results['Keywords'], list):
        return "Public", "Không tìm thấy từ khóa hợp lệ trong kết quả."

    # Extract keywords and patterns from results
    try:
        found_keywords = {item['Found Keyword'] for item in results['Keywords'] if isinstance(
            item, dict) and 'Found Keyword' in item}
        found_patterns = {item['Pattern Name']: item['num of the same pattern']
                          for item in results['Patterns'] if isinstance(item, dict) and 'Pattern Name' in item and 'num of the same pattern' in item}
    except TypeError as e:
        return "Internal", f"Đã xảy ra lỗi khi phân tích kết quả: {str(e)}"

    # Helper function to check if the result matches the rules
    def matches_keywords_rule(result_keywords, rule_keywords):
        if isinstance(rule_keywords, list):
            return any(keyword in result_keywords for keyword in rule_keywords)
        else:
            return rule_keywords in result_keywords

    # Rule 1: Match at least 3 values and check if certain patterns have count >= 10
    rule_1_matches = 0
    matched_rule_1_keys = []
    matched_pattern_counts = []

    for key, value in rules['rule_1'].items():
        # Check if either keyword or pattern matches
        keyword_match = matches_keywords_rule(found_keywords, value)
        # Ensure pattern count is >= 10
        pattern_match = key in found_patterns and found_patterns[key] >= 10

        if keyword_match or pattern_match:
            rule_1_matches += 1
            matched_rule_1_keys.append(key)
            # Add the number of occurrences of the pattern to the list
            if pattern_match:
                matched_pattern_counts.append(f"{key}: {found_patterns[key]}")

    # If at least 3 keys match, assign "Confidential"
    if rule_1_matches >= 3:
        # Include pattern counts in the sms_scan message
        sms_scan = f"Tìm thấy nội dung mật rule 1: {rule_1_matches} key(s) matched: {', '.join(matched_rule_1_keys)}. Pattern counts: {', '.join(matched_pattern_counts)}"
        pretty_sms_scan = convert_string_to_json_rule_one(sms_scan)
        return "Confidential", pretty_sms_scan

    # Rule 2: Check both keyword and pattern match for all keys in rule_2
    rule_2_matches = True
    matched_rule_2_keys = []

    for key, value in rules['rule_2'].items():
        keyword_match = matches_keywords_rule(
            found_keywords, value)  # Check if keyword matches
        pattern_match = key in found_patterns  # Check if pattern name matches

        # If neither keyword nor pattern matches, rule_2 does not match
        if not (keyword_match or pattern_match):
            rule_2_matches = False
            break
        matched_rule_2_keys.append(key)

    if rule_2_matches:
        sms_scan = f"Tìm thấy nội dung mật rule 2: {', '.join(matched_rule_2_keys)}"
        pretty_sms_scan = convert_string_to_json_rule_two(sms_scan)
        return "Confidential", pretty_sms_scan

    # Rule 3: At least 6 values must match
    rule_3_matches = 0
    matched_rule_3_keys = []

    for key, value in rules['rule_3'].items():
        if matches_keywords_rule(found_keywords, value):
            rule_3_matches += 1
            matched_rule_3_keys.append(key)

    if rule_3_matches >= 6:
        sms_scan = f"Tìm thấy nội dung mật rule 3: {', '.join(matched_rule_3_keys)}"
        return "Confidential", sms_scan

    # Rule 4: At least one keyword must match
    matched_rule_4_keywords = [
        keyword for keyword in rules['rule_4'] if keyword in found_keywords]

    if matched_rule_4_keywords:
        sms_scan = f"Tìm thấy nội dung mật rule 3: {', '.join(matched_rule_4_keywords)}"
        pretty_sms_scan = convert_rule_string_to_json_rule_four(sms_scan)
        return "Confidential", pretty_sms_scan

    # If no conditions match, return "Internal"
    return "Internal", "Không tìm thấy nội dung mật"


# Ví dụ sử dụng
# file_path = "/Users/vuanhtien/Documents/CodeForMe/BIDV/label_tool/Chi test 4.docx"

# results = scan_file(file_path)
# print("results:", type(results))
# print("results:", results)

# results2 = scan_file(file_path)
# rul = define_rules()
# mapping2 = classify_document_with_multiple_rules(results2, rul)

# print("kết quả phân loại:", mapping2)


def label_docx_file(file_path, classify):
    """Add or overwrite a label in the footer of a DOCX file based on the scan results."""
    document = Document(file_path)

    scan_result = scan_file(file_path)
    rul = define_rules()
    classify, sms = classify_document_with_multiple_rules(scan_result, rul)
    # Determine the label text based on the category
    if classify == "Confidential":
        label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    elif classify == "Internal":
        label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"

    section = document.sections[0]

    # Ensure layout margins are correct
    section.top_margin = Cm(2)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    footer = section.footer

    # Remove all existing paragraphs in the footer
    for element in footer._element.xpath('.//w:p'):
        element.getparent().remove(element)

    # Add a new paragraph that contains both the page number and the label text
    paragraph = footer.add_paragraph()

    # Set a tab stop at the right margin for the page number
    right_margin_position = section.page_width - \
        section.right_margin - section.left_margin
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        right_margin_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)

    # Add the label text (left-aligned)
    run_label = paragraph.add_run(label_text)
    run_label.font.name = 'Times New Roman'
    run_label.font.size = Pt(12)

    # Add a tab to move the cursor to the right side of the page
    run_label.add_tab()

    # Add the page number field (aligned to the right tab stop)
    run_page = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')  # Begin field
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')  # Field instruction text
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')  # Separate field
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')  # End field
    fldChar3.set(qn('w:fldCharType'), 'end')

    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    run_page._r.append(fldChar3)

    # Set font for the page number
    run_page.font.name = 'Times New Roman'
    run_page.font.size = Pt(12)

    # Align the paragraph to left (default behavior)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document
    document.save(file_path)

    return label_text


# def label_xlsx_file_watermark(file_path, category, image_path_public, image_path_confidential):
#     """Add watermark to an Excel file based on scan results."""
#     workbook = load_workbook(file_path)
#     sheet = workbook.active

#     if category == "Confidential":
#         label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
#         img_path = image_path_confidential
#     else:
#         label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
#         img_path = image_path_public

#     img = Image(img_path)
#     img.width, img.height = 400, 200
#     sheet.add_image(img, 'A1')

#     workbook.save(file_path)

#     return label_text


# Function to label XLSX file footer


def label_xlsx_file_footer(file_path, classify):
    """Add footer label to an Excel file based on scan results."""

    # Scan the file to determine the classification
    scan_result = scan_file(file_path)
    rul = define_rules()
    classify, sms = classify_document_with_multiple_rules(scan_result, rul)

    # Determine the label text based on the classification
    if classify == "Confidential":
        label_text = "Tài liệu Mật của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"
    elif classify == "Internal":
        label_text = "Tài liệu Nội bộ của BIDV. Cấm sao chép, in ấn dưới mọi hình thức!"


    # Load the Excel workbook
    workbook = load_workbook(file_path)

    # Add label to the footer of each worksheet
    for sheet in workbook.worksheets:
        # Construct the new footer with label on the left and page number on the right
        new_footer = f"&L{label_text} &RTrang &P"

        # Set the new footer for the sheet
        sheet.oddFooter.center.text = new_footer
        sheet.oddFooter.center.size = 12
        sheet.oddFooter.center.font = "Times New Roman"

    # Save the modified workbook
    workbook.save(file_path)

    return label_text
